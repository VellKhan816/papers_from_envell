import os
from datetime import datetime
from flask import Flask, render_template, request, redirect, url_for, flash, session, jsonify, send_file

app = Flask(__name__)
app.secret_key = 'your-super-secret-key-change-in-production'

DOCUMENTS_DIR = os.path.join(os.path.dirname(__file__), 'documents')
TRASH_DIR = os.path.join(os.path.dirname(__file__), 'trash')
os.makedirs(DOCUMENTS_DIR, exist_ok=True)
os.makedirs(TRASH_DIR, exist_ok=True)

# Тестовые пользователи
USERS = {
    "admin": "admin123",
    "user": "user321"
}


def get_documents(folder, is_trash=False):
    docs = []
    for filename in os.listdir(folder):
        if filename.endswith('.txt'):
            filepath = os.path.join(folder, filename)
            if is_trash:
                # Удаляем файлы старше 30 дней
                file_age = datetime.now() - datetime.fromtimestamp(os.path.getmtime(filepath))
                if file_age.days > 30:
                    for ext in ['.txt', '.docx', '.pdf']:
                        try:
                            os.remove(filepath.replace('.txt', ext))
                        except:
                            pass
                    continue

            with open(filepath, 'r', encoding='utf-8') as f:
                lines = f.read().split('\n', 3)
                title = lines[0].replace('Заголовок: ', '') if len(lines) > 0 else 'Без названия'
                date_line = lines[1] if len(lines) > 1 else ''
                date = date_line.replace('Дата: ', '') if date_line else 'Неизвестно'
                preview = lines[3][:100] + '...' if len(lines) > 3 else 'Пустой документ'
            docs.append({
                "filename": filename,
                "title": title,
                "date": date,
                "preview": preview
            })
    return sorted(docs, key=lambda x: x['filename'], reverse=True)


def count_trash_files():
    count = 0
    for filename in os.listdir(TRASH_DIR):
        if filename.endswith('.txt'):
            filepath = os.path.join(TRASH_DIR, filename)
            file_age = datetime.now() - datetime.fromtimestamp(os.path.getmtime(filepath))
            if file_age.days <= 30:
                count += 1
            else:
                for ext in ['.txt', '.docx', '.pdf']:
                    try:
                        os.remove(filepath.replace('.txt', ext))
                    except:
                        pass
    return count


@app.route('/')
def welcome():
    return render_template('welcome.html')


@app.route('/login', methods=['GET', 'POST'])
def login():
    if request.method == 'POST':
        login = request.form['login']
        password = request.form['password']
        if login in USERS and USERS[login] == password:
            session['user'] = login
            flash("Вход выполнен успешно!", "success")
            return redirect(url_for('dashboard'))
        else:
            flash("Неверный логин или пароль", "error")
    return render_template('login.html')


@app.route('/dashboard')
def dashboard():
    if 'user' not in session:
        return redirect(url_for('login'))
    documents = get_documents(DOCUMENTS_DIR)
    trash_count = count_trash_files()
    return render_template('dashboard.html', user=session['user'], documents=documents, page='dashboard',
                           trash_count=trash_count)


@app.route('/trash')
def trash():
    if 'user' not in session:
        return redirect(url_for('login'))
    documents = get_documents(TRASH_DIR, is_trash=True)
    trash_count = len(documents)
    return render_template('trash.html', user=session['user'], documents=documents, page='trash',
                           trash_count=trash_count)


@app.route('/editor')
def editor():
    if 'user' not in session:
        return redirect(url_for('login'))
    doc_file = request.args.get('doc')
    title = 'Без названия'
    content = ''
    folder = DOCUMENTS_DIR
    if doc_file and doc_file.endswith('.txt'):
        if os.path.exists(os.path.join(DOCUMENTS_DIR, doc_file)):
            folder = DOCUMENTS_DIR
        elif os.path.exists(os.path.join(TRASH_DIR, doc_file)):
            folder = TRASH_DIR
        else:
            return "Документ не найден", 404

        filepath = os.path.join(folder, doc_file)
        with open(filepath, 'r', encoding='utf-8') as f:
            lines = f.read().split('\n', 3)
            title = lines[0].replace('Заголовок: ', '') if len(lines) > 0 else 'Без названия'
            content = lines[3] if len(lines) > 3 else ''
    return render_template('editor.html', title=title, content=content, doc_file=doc_file, folder=folder)


@app.route('/save-document', methods=['POST'])
def save_document():
    if 'user' not in session:
        return jsonify({"error": "Не авторизован"}), 401

    title = request.form.get('title', 'Без названия')
    content = request.form.get('content', '')
    doc_file = request.form.get('doc_file')
    folder = request.form.get('folder', DOCUMENTS_DIR)

    if not doc_file:
        timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
        doc_file = f"document_{timestamp}.txt"

    filepath = os.path.join(folder, doc_file)
    with open(filepath, 'w', encoding='utf-8') as f:
        f.write(f"Заголовок: {title}\n")
        f.write(f"Дата: {datetime.now().strftime('%d %b. %Y г.')}\n")
        f.write("=" * 40 + "\n\n")
        f.write(content)

    # DOCX
    try:
        from docx import Document
        docx_path = filepath.replace('.txt', '.docx')
        doc = Document()
        doc.add_heading(title, 0)
        doc.add_paragraph(content)
        doc.save(docx_path)
    except:
        pass

    # PDF
    try:
        from fpdf import FPDF
        pdf_path = filepath.replace('.txt', '.pdf')
        pdf = FPDF()
        pdf.add_page()
        pdf.set_font('Arial', 'B', 16)
        pdf.cell(0, 10, title, 0, 1, 'C')
        pdf.ln(10)
        pdf.set_font('Arial', '', 12)
        pdf.multi_cell(0, 10, content)
        pdf.output(pdf_path)
    except:
        pass

    return jsonify({"success": True, "filename": doc_file, "message": f"✅ Документ '{title}' сохранён!"})


@app.route('/delete-document/<filename>', methods=['POST'])
def delete_document(filename):
    if 'user' not in session:
        return jsonify({"error": "Не авторизован"}), 401

    for ext in ['.txt', '.docx', '.pdf']:
        src = os.path.join(DOCUMENTS_DIR, filename.replace('.txt', ext))
        dst = os.path.join(TRASH_DIR, filename.replace('.txt', ext))
        if os.path.exists(src):
            os.rename(src, dst)

    return jsonify({"success": True, "message": "Документ перемещён в корзину"})


@app.route('/restore-document/<filename>', methods=['POST'])
def restore_document(filename):
    if 'user' not in session:
        return jsonify({"error": "Не авторизован"}), 401

    for ext in ['.txt', '.docx', '.pdf']:
        src = os.path.join(TRASH_DIR, filename.replace('.txt', ext))
        dst = os.path.join(DOCUMENTS_DIR, filename.replace('.txt', ext))
        if os.path.exists(src):
            os.rename(src, dst)

    return jsonify({"success": True, "message": "Документ восстановлен"})


@app.route('/empty-trash', methods=['POST'])
def empty_trash():
    if 'user' not in session:
        return jsonify({"error": "Не авторизован"}), 401

    for filename in os.listdir(TRASH_DIR):
        if any(filename.endswith(ext) for ext in ['.txt', '.docx', '.pdf']):
            try:
                os.remove(os.path.join(TRASH_DIR, filename))
            except:
                pass

    return jsonify({"success": True, "message": "Корзина очищена"})


@app.route('/download/<format>/<filename>')
def download_document(format, filename):
    if format not in ['txt', 'docx', 'pdf']:
        return "Неверный формат", 400
    for folder in [DOCUMENTS_DIR, TRASH_DIR]:
        filepath = os.path.join(folder, filename.replace('.txt', f'.{format}'))
        if os.path.exists(filepath):
            return send_file(filepath, as_attachment=True)
    return "Файл не найден", 404


@app.route('/logout')
def logout():
    session.pop('user', None)
    return redirect(url_for('welcome'))


if __name__ == '__main__':
    app.run(debug=True)